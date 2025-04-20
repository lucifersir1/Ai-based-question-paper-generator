from flask import Flask, render_template, request, redirect, url_for, session, send_file
from auth import auth_bp
from model import model  # Import the Gemini model from model.py
import sqlite3
import os
import uuid
from werkzeug.utils import secure_filename
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

app = Flask(__name__)
app.secret_key = 'R31R32S33A350000000000'

# Register auth blueprint
app.register_blueprint(auth_bp)

UPLOAD_FOLDER = 'uploads'
GENERATED_FOLDER = 'generated_papers'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(GENERATED_FOLDER, exist_ok=True)
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['GENERATED_FOLDER'] = GENERATED_FOLDER

ALLOWED_EXTENSIONS = {'txt', 'pdf', 'docx'}

# def update_question_papers_table():
#     conn = sqlite3.connect('database.db')
#     cursor = conn.cursor()

#     cursor.execute("PRAGMA table_info(question_papers)")
#     columns = [col[1] for col in cursor.fetchall()]
#     if 'teacher_name' not in columns:
#         cursor.execute("ALTER TABLE question_papers ADD COLUMN teacher_name TEXT")
#         conn.commit()

#     conn.close()

# # Call the function to ensure schema is updated
# update_question_papers_table()


# Utility function to check allowed file types
def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

# Generate questions using Gemini API from model.py
def generate_questions(subject_code, difficulty, syllabus_path, num_questions):
    """Generate AI-based questions using the Gemini model."""
    with open(syllabus_path, 'r', encoding='utf-8') as file:
        syllabus_content = file.read()

    prompt = f"""
    Based on the syllabus below, generate {num_questions} questions for the subject {subject_code}.
    Difficulty: {difficulty}
    Syllabus:
    {syllabus_content}
    """

    print("AI Prompt Sent: ", prompt)

    # Use the Gemini model to generate questions
    try:
        response = model.generate_content(prompt)
        ai_generated_output = response.text.split("\n")

        # Process output to extract questions
        ai_questions = [
            {"text": q.strip(), "marks": 5}
            for q in ai_generated_output if q.strip()
        ][:num_questions]

        print("Generated AI Questions: ", ai_questions)
        return ai_questions, prompt

    except Exception as e:
        print(f"❌ Error generating questions: {e}")
        return [{"text": "Error generating questions.", "marks": 5}], prompt

# Inserts questions dynamically into DOCX template


def insert_questions_into_docx(template_path, questions, output_path):
    doc = Document(template_path)

    # Find and replace {{QUESTIONS}} placeholder
    for paragraph in doc.paragraphs:
        if "{{QUESTIONS}}" in paragraph.text:
            paragraph.clear()  # Clear the placeholder

            # Insert questions with customized styling
            for idx, main_q in enumerate(questions, start=1):
                # Main Question Formatting
                main_question = f"Q{idx}. {main_q['heading']} [{main_q['marks']} Marks]"
                main_run = paragraph.add_run(main_question)
                paragraph.add_run("\n")
                # Set font style and size for Main Question
                main_run.font.name = 'Times New Roman'
                main_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
                main_run.font.size = Pt(13)
                main_run.font.bold = True

                paragraph.add_run("\n")  # Add new line

                # Sub-Questions Formatting
                for sub_idx, sub_q in enumerate(main_q['sub_questions'], start=1):
                    sub_question = f"    ({chr(96 + sub_idx)}) {sub_q['text']} [{sub_q['marks']} Marks]"
                    sub_run = paragraph.add_run(sub_question)

                    # Set font style and size for Sub-Questions
                    main_run.font.name = 'Times New Roman'
                    main_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
                    main_run.font.size = Pt(13)
                    main_run.font.bold = True
                    paragraph.add_run("\n")  # Add new line after each sub-question

            break  # Stop after replacing the first occurrence

    # Save the updated document
    doc.save(output_path)
    # # Prepare questions in formatted text
    # question_text = ""
    # for idx, main_q in enumerate(questions, start=1):
    #     question_text += f"Q{idx}. {main_q['heading']} [{main_q['marks']} Marks]\n"
    #     for sub_idx, sub_q in enumerate(main_q['sub_questions'], start=1):
    #         question_text += f"    ({chr(96 + sub_idx)}) {sub_q['text']} [{sub_q['marks']} Marks]\n"

    # # Replace the {{QUESTIONS}} placeholder
    # for paragraph in doc.paragraphs:
    #     if "{{QUESTIONS}}" in paragraph.text:
    #         paragraph.text = paragraph.text.replace("{{QUESTIONS}}", question_text)

    # # Save the updated document
    # doc.save(output_path)


@app.route('/teacher_dashboard', methods=['GET', 'POST'])
def teacher_dashboard():
    if 'user_id' not in session or session.get('role') != 'teacher':
        return redirect(url_for('auth.login'))

    conn = sqlite3.connect('database.db')
    cursor = conn.cursor()
    cursor.execute("SELECT teacher_name, subject FROM users WHERE id = ?", (session['user_id'],))
    result = cursor.fetchone()
    teacher_name, subject = result if result else ("Teacher", "Unknown")
    conn.close()

    if request.method == 'POST':
        # Collect form data
        subject_code = request.form['subject_code']
        difficulty = request.form['difficulty']

        # Handle syllabus and paper format uploads
        syllabus_file = request.files.get('syllabus_file')
        paper_format_file = request.files.get('paper_format_file')

        if not syllabus_file or not allowed_file(syllabus_file.filename) or \
           not paper_format_file or not allowed_file(paper_format_file.filename):
            return "Invalid or missing file. Only .txt, .pdf, and .docx are allowed.", 400

        # Save uploaded files securely
        syllabus_filename = secure_filename(f"{uuid.uuid4()}_{syllabus_file.filename}")
        paper_format_filename = secure_filename(f"{uuid.uuid4()}_{paper_format_file.filename}")

        syllabus_path = os.path.join(app.config['UPLOAD_FOLDER'], syllabus_filename)
        paper_format_path = os.path.join(app.config['UPLOAD_FOLDER'], paper_format_filename)

        syllabus_file.save(syllabus_path)
        paper_format_file.save(paper_format_path)

        # Parse dynamic question input
        num_main_questions = int(request.form['num_main_questions'])
        questions = []

        for i in range(1, num_main_questions + 1):
            main_heading = request.form.get(f'main_question_{i}_text')
            main_marks = int(request.form.get(f'main_question_{i}_marks'))
            num_sub_questions = int(request.form.get(f'num_sub_questions_{i}'))

            # Generate sub-questions dynamically with AI
            sub_questions = []
            ai_generated_subs, prompt = generate_questions(subject_code, difficulty, syllabus_path, num_sub_questions)

            for j in range(num_sub_questions):
                sub_marks = int(request.form.get(f'main_question_{i}_sub_{j + 1}_marks'))
                sub_text = ai_generated_subs[j]['text'] if j < len(ai_generated_subs) else "AI-generated question"
                sub_questions.append({'text': sub_text, 'marks': sub_marks})

            questions.append({
                'heading': main_heading,
                'marks': main_marks,
                'sub_questions': sub_questions
            })

        # Generate the output DOCX file
        output_filename = f"generated_{uuid.uuid4()}.docx"
        output_path = os.path.join(app.config['GENERATED_FOLDER'], output_filename)

        insert_questions_into_docx(paper_format_path, questions, output_path)

        conn = sqlite3.connect('database.db')
        cursor = conn.cursor()
        cursor.execute('''
            INSERT INTO question_papers (subject, subject_code, teacher_name, difficulty)
            VALUES (?, ?, ?, ?)
        ''', (subject, subject_code, teacher_name, difficulty))

        conn.commit()
        conn.close()
        # Clean up temporary uploaded files
        os.remove(syllabus_path)
        os.remove(paper_format_path)

        # Render the result page with download link
        return render_template(
            'result.html',
            subject_code=subject_code,
            questions=questions,
            total_marks=sum(q['marks'] for q in questions),
            download_file=output_filename,
            prompt=prompt
        )

    return render_template('teacher_dashboard.html', teacher_name=teacher_name)

@app.route('/admin_dashboard', methods=['GET'])
def admin_dashboard():
    # Ensure only admin can access
    if 'user_id' not in session or session.get('role') != 'admin':
        return redirect(url_for('auth.login'))

    # Connect to SQLite
    conn = sqlite3.connect('database.db')
    cursor = conn.cursor()

    # Fetch distinct subjects (where role is 'teacher')
    cursor.execute("SELECT DISTINCT subject FROM users WHERE role='teacher'")
    subjects = [row[0] for row in cursor.fetchall()]
    
    print("Available Subjects:", subjects)  # Debugging Output

    # Get the selected subject from the form
    selected_subject = request.args.get('subject')
    print("Selected Subject:", selected_subject)  # Debugging Output

    # Fetch question papers (optionally filter by subject)
    if selected_subject:
        cursor.execute('''
            SELECT teacher_name, difficulty
            FROM question_papers
            WHERE subject = ?
        ''', (selected_subject,))
    else:
        cursor.execute('''
            SELECT teacher_name, difficulty
            FROM question_papers
        ''')
    # Store data in a list of dictionaries
    question_papers = [
        {
            'teacher_name': row[0],
            'difficulty': row[1]
        }
        for row in cursor.fetchall()
    ]

    conn.close()

    return render_template('admin_dashboard.html',
                           subjects=subjects,
                           question_papers=question_papers,
                           selected_subject=selected_subject)



@app.route('/download/<int:paper_id>')
def download_paper(paper_id):
    conn = sqlite3.connect('database.db')
    cursor = conn.cursor()
    cursor.execute('SELECT output_path FROM question_papers WHERE id = ?', (paper_id,))
    paper = cursor.fetchone()
    conn.close()

    if paper:
        return send_file(paper[0], as_attachment=True)
    else:
        return "File not found", 404


@app.route('/download/<filename>')
def download(filename):
    file_path = os.path.join(app.config['GENERATED_FOLDER'], filename)
    if os.path.exists(file_path) and allowed_file(filename):
        return send_file(file_path, as_attachment=True)
    return "File not found or invalid.", 404

if __name__ == '__main__':
    app.run(debug=True)
